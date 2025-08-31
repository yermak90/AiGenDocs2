from docx import Document

def build_contract(context: dict, out_path: str):
    doc = Document()
    doc.add_heading("ДОГОВОР № " + context["number"])
    doc.add_paragraph("Стороны: " + context["company"] + " и " + context["counterparty"])
    doc.save(out_path)
